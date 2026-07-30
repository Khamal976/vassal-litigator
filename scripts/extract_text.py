#!/usr/bin/env python3
"""
extract_text.py — детерминированная Python-нога OCR-модуля `shared/ocr`.

Контракт, дерево решений и агентский протокол (vision, кэш, структурный режим,
пер-полевой confidence) — см. shared/ocr.md. Эта программа делает ТОЛЬКО
детерминированную часть и НЕ выполняет vision (Python не имеет доступа к Read tool):

  - программное извлечение текста: PDF с текстовым слоем, DOCX, TXT;
  - структурная сводка таблиц (G.3): `.xlsx`/`.xlsm`/крупный `.csv` → скелет в зеркало
    (листы, шапки, образец строк, итоговые строки, диапазон дат, скрытые листы,
    детект формул без сохранённых значений). Приём таблицу НЕ анализирует — суммы,
    сверки и флаги считает `analyze_table.py` на стороне `study-evidence`;
  - детект «мусорного» текстового слоя (F3.3) — mojibake / тонкий OCR-артефакт
    поверх скана → директива needs_vision вместо молчаливого возврата мусора;
  - решение needs_vision / vision_pages / vision_reason (F3.1): vision — основной
    путь OCR, не tesseract-rus (которого в Cowork нет и который упирается в таймаут);
  - рекомендация структурного режима для длинных сканов (F3.4);
  - вычисление content_hash для кэша (F3.2; само решение о кэше — агентское);
  - рендер страницы(ц) в PNG для vision через pymupdf (надёжнее системного
    pdftoppm, который у файлового Read падает «unsafe location») + guard обрезки (F3.5);
  - tesseract спот-сверка критичного поля по вендоренному rus.traineddata (F3.1,
    вторая нога) — независимое второе чтение цифры (ИНН/сумма/доля/дата) с машинным
    вердиктом match/mismatch/inconclusive (F.18). Деградирует тихо: нет словаря или
    бинаря, таймаут → available:false, приём продолжается на результате vision.

Режимы запуска:
    # 1) извлечение (основной режим, вызывается скиллами ingest):
    python3 extract_text.py <файл> [--output-dir <папка>] [--render-dir <папка>]
                            [--max-head-pages N] [--structural-threshold N]

    # 2) рендер одной страницы в PNG (для vision-ноги агента):
    python3 extract_text.py --render <file.pdf> --page N [--render-dir <папка>] [--width 1500]

    # 3) tesseract спот-сверка критичного поля (второе, независимое чтение):
    python3 extract_text.py --spot-check <file.pdf> --page N --expect "3390000"
                            [--region x0,y0,x1,y1]   # доли страницы, если целая — таймаут
                            [--psm N] [--tessdata <путь>] [--render-dir <папка>]
    python3 extract_text.py --spot-check <crop.png> --psm 7 --expect "7707083893"

Выход: JSON. Поля контракта — см. shared/ocr.md §3.
"""

import sys
import json
import os
import re
import errno
import hashlib
import subprocess
import shutil
import glob
from pathlib import Path

# --- Параметры по умолчанию -------------------------------------------------

DEFAULT_RENDER_WIDTH = 1500          # ориентир ширины PNG для vision (1200–1700 px)
DEFAULT_MAX_HEAD_PAGES = 10          # сколько головных страниц полнотекстить в структурном режиме
DEFAULT_STRUCTURAL_THRESHOLD = 15    # порог (скан-страниц) для рекомендации структурного режима


# --- content_hash (F3.2) ----------------------------------------------------

def compute_content_hash(filepath: str) -> str:
    """SHA-256 от бинарного содержимого файла (ключ кэша/дедупа)."""
    h = hashlib.sha256()
    try:
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(1 << 20), b""):
                h.update(chunk)
        return "sha256:" + h.hexdigest()
    except Exception:
        return ""


# --- Детект мусорного текстового слоя (F3.3) --------------------------------

def _text_diagnostics(text: str) -> dict:
    """Метрики качества извлечённого текстового слоя страницы/документа."""
    t = text or ""
    total = len(t)
    if total == 0:
        return {"empty": True, "len": 0}
    repl = t.count("�")
    ctrl = sum(1 for ch in t if ord(ch) < 32 and ch not in "\t\n\r\f")
    alpha = [ch for ch in t if ch.isalpha()]
    cyr = sum(1 for ch in alpha if "Ѐ" <= ch <= "ӿ")
    cyr_ratio = (cyr / len(alpha)) if alpha else 0.0
    tokens = [w for w in re.split(r"\s+", t.strip()) if w]
    avg_tok = (sum(len(w) for w in tokens) / len(tokens)) if tokens else 0.0
    return {
        "empty": False,
        "len": total,
        "repl_ratio": repl / total,
        "ctrl_ratio": ctrl / total,
        "cyr_ratio": cyr_ratio,
        "alpha_count": len(alpha),
        "avg_token_len": avg_tok,
    }


def _is_garbage(diag: dict) -> bool:
    """True, если текстовый слой — мусор (mojibake/каша), а не осмысленный текст.

    Пусто (empty) — это НЕ мусор: это «нет текстового слоя» (скан) → отдельная ветка.
    Пороги консервативны, чтобы не гнать на vision годный текст.
    """
    if diag.get("empty"):
        return False
    if diag.get("repl_ratio", 0) > 0.02:          # символы замены — сильный сигнал порчи кодировки
        return True
    if diag.get("ctrl_ratio", 0) > 0.05:          # много управляющих символов
        return True
    # текста достаточно, но кириллицы почти нет — характерно для cp1251-mojibake рус. документа.
    # Порог 30 букв: на реальной странице мусора их сотни; ложное срабатывание лишь отправит
    # годную (напр. целиком латинскую) страницу на vision — vision её корректно транскрибирует.
    if diag.get("alpha_count", 0) >= 30 and diag.get("cyr_ratio", 1.0) < 0.15:
        return True
    # «словесная каша»: неправдоподобно короткие токены при объёме
    if diag.get("len", 0) > 200 and diag.get("avg_token_len", 99) < 2.0:
        return True
    return False


def _classify_page(page) -> str:
    """Статус страницы PDF: 'text' | 'garbage' | 'scan' | 'blank'."""
    txt = page.get_text()
    stripped = txt.strip()
    if len(stripped) < 10:
        # почти нет текста — скан, если есть изображения; иначе пустая страница
        try:
            has_images = bool(page.get_images())
        except Exception:
            has_images = False
        return "scan" if has_images else "blank"
    if _is_garbage(_text_diagnostics(txt)):
        return "garbage"
    return "text"


# --- Детект составного пакета (E.6.2) ---------------------------------------

# Заголовки-якоря, которые в норме СТОЯТ В ШАПКЕ самостоятельного документа
# (не встречаются как обычное слово в теле). Повтор одного якоря в шапке ≥2 страниц —
# сигнал, что в одном PDF склеены несколько документов одного типа (напр. 2 платёжки).
# Матчинг — по нормализованной (ё→е, нижний регистр, схлопнутые пробелы) шапке страницы.
_COMPOSITE_HEADER_ANCHORS = [
    "платежное поручение",
    "счет-фактура",
    "счет на оплату",
    "универсальный передаточный документ",
    "товарная накладная",
    "товарно-транспортная накладная",
    "приходный кассовый ордер",
    "расходный кассовый ордер",
]
# Судебные акты: якорь засчитывается ТОЛЬКО вместе с «арбитражный суд» в шапке —
# иначе «определение/решение» ловится в теле («суд вынес определение …»).
_COMPOSITE_COURT_ANCHORS = ["определение", "решение", "постановление"]

_COMPOSITE_HEAD_CHARS = 300          # «шапка» страницы — первые N символов
_INN_RE = re.compile(r"\b\d{10}\b|\b\d{12}\b")


def _norm_head(text: str) -> str:
    """Нормализованная шапка страницы для матчинга якорей."""
    head = (text or "")[:_COMPOSITE_HEAD_CHARS].lower().replace("ё", "е")
    return re.sub(r"\s+", " ", head)


def _detect_composite(text_parts, statuses) -> dict:
    """Эвристика склейки нескольких самостоятельных документов в одном PDF (E.6.2).

    Консервативно, чтобы НЕ гнать иск-с-приложениями (у него один заголовок, а
    приложения — с разными шапками): триггер — один заголовок-якорь в шапке ≥2
    РАЗНЫХ страниц. Множественные ИНН — только вспомогательная пометка при уже
    сработавшем триггере. Решение о split — агентское (в preview), здесь директива.
    Скан/мусор-страницы пропускаем: их текста ещё нет (vision отработает позже).
    """
    anchor_pages = {}   # ярлык якоря -> множество индексов страниц
    for i, (part, status) in enumerate(zip(text_parts, statuses)):
        if status != "text":
            continue
        head = _norm_head(part)
        for anchor in _COMPOSITE_HEADER_ANCHORS:
            if anchor in head:
                anchor_pages.setdefault(anchor, set()).add(i)
        if "арбитражный суд" in head:
            for anchor in _COMPOSITE_COURT_ANCHORS:
                if anchor in head:
                    anchor_pages.setdefault(anchor, set()).add(i)

    reasons = []
    for anchor, pages in sorted(anchor_pages.items()):
        if len(pages) >= 2:
            nums = ", ".join(str(p + 1) for p in sorted(pages))
            reasons.append(f"заголовок «{anchor}» — в шапке {len(pages)} стр. ({nums})")

    if reasons:
        inns = set()
        for part, status in zip(text_parts, statuses):
            if status == "text":
                inns.update(_INN_RE.findall(part or ""))
        if len(inns) >= 2:
            reasons.append(f"различных ИНН в документе: {len(inns)}")

    return {"composite_suspected": bool(reasons), "composite_reasons": reasons}


# --- Классификация сбоя открытия файла (E.14 а/б) ---------------------------

# Маркеры реальной порчи PDF в тексте исключения pymupdf/mupdf.
_CORRUPT_MARKERS = (
    "format error", "cannot open broken", "broken document", "no objects found",
    "syntax error", "damaged", "not a pdf", "cannot recognize", "bad xref",
)
# errno, характерные для НЕ материализованного файла на облачном маунте (OneDrive).
_MATERIALIZE_ERRNOS = {errno.ENOENT, errno.EIO, errno.ESTALE, errno.EACCES, errno.EAGAIN}
# Текстовые маркеры того же (когда errno недоступен — pymupdf оборачивает по-своему).
_MATERIALIZE_MSG_MARKERS = (
    "no such file", "cannot find", "not materialized", "input/output error", "stale file",
)


def classify_open_error(filepath: str, exc: Exception) -> tuple:
    """Различает 'файл не материализован (OneDrive/маунт)' vs 'реально битый' (E.14 а/б).

    Возвращает (error_class, retryable, message):
      - not_materialized / retryable=True  — файл-заглушка не синкнут; ретрай с паузой (rule 2).
      - corrupt          / retryable=False — данные PDF повреждены; ветка «битый файл».
      - empty            / retryable=False — файл нулевого размера.
    Порядок проверок: пусто → не материализован → битый → (по умолчанию) битый.
    """
    exc_name = type(exc).__name__
    msg = str(exc).lower()

    # 0-байтовый файл (в т.ч. pymupdf EmptyFileError) — не битый и не «дождись синка».
    try:
        if os.path.getsize(filepath) == 0:
            return ("empty", False, "файл нулевого размера")
    except OSError:
        pass  # getsize сам упал → отнесём к not_materialized ниже
    if exc_name == "EmptyFileError":
        return ("empty", False, "файл нулевого размера")

    # I/O-ошибки и отсутствие файла на маунте → вероятная дегидратация OneDrive → ретрай.
    oserrno = getattr(exc, "errno", None)
    if (isinstance(exc, FileNotFoundError) or oserrno in _MATERIALIZE_ERRNOS
            or any(m in msg for m in _MATERIALIZE_MSG_MARKERS)):
        return ("not_materialized", True,
                "файл не открылся на маунте — возможно, не синхронизирован OneDrive "
                "(cloud-only): сначала Read (материализует), затем ретрай с паузой")

    # Явные маркеры порчи PDF (pymupdf FileDataError и т.п.) → реально битый.
    if exc_name == "FileDataError" or any(m in msg for m in _CORRUPT_MARKERS):
        return ("corrupt", False, f"PDF повреждён, не открывается: {exc}")

    # Неизвестная ошибка на существующем ненулевом файле — считаем битым (не ретраить бесконечно).
    return ("corrupt", False, f"PDF не открывается: {exc}")


def _pdf_error(error_class: str, retryable: bool, message: str, pages: int = 0) -> dict:
    """Единый JSON-отказ PDF-извлечения с классификацией сбоя (E.14)."""
    return {"text": "", "method": "none", "confidence": "low", "pages": pages,
            "page_statuses": [], "needs_vision": False, "vision_pages": [],
            "vision_pages_suggested": [], "vision_reason": None,
            "structural_recommended": False,
            "composite_suspected": False, "composite_reasons": [],
            "error_class": error_class, "retryable": retryable,
            "warnings": [message]}


# --- Движок poppler: fallback, когда нет pymupdf (F.2) ----------------------
#
# Зачем. pymupdf был ЕДИНСТВЕННЫМ путём к PDF: без него не работало ни
# программное извлечение, ни рендер страниц — а значит и vision, потому что
# vision читает PNG, который рендерил тот же pymupdf. Задокументированный в
# ocr.md fallback «нет pymupdf → vision» был замкнут сам на себя. В боевых
# прогонах pymupdf регулярно не вставал (25-МБ колесо, прокси, а до F.1 —
# ещё и падавший setup.sh), и агент каждый раз вручную уходил на poppler.
# Здесь этот обход сделан штатным: poppler (pdftotext/pdftoppm/pdfinfo)
# предустановлен в песочнице Cowork.
#
# Граница движка. poppler не отдаёт список изображений страницы, поэтому
# страницу без текстового слоя нельзя отличить от пустой. Классифицируем её
# как `scan` (консервативно: лишняя страница уйдёт на vision — это дешевле,
# чем принять скан за пустую и потерять содержание).

def _poppler_bin(name: str):
    """Путь к утилите poppler или None."""
    return shutil.which(name)


def poppler_available() -> bool:
    return bool(_poppler_bin("pdftotext"))


def _pdf_page_count_poppler(filepath: str) -> int:
    """Число страниц через pdfinfo; 0 — если не удалось."""
    exe = _poppler_bin("pdfinfo")
    if not exe:
        return 0
    try:
        r = subprocess.run([exe, filepath], capture_output=True, text=True, timeout=60)
    except Exception:
        return 0
    m = re.search(r"^Pages:\s+(\d+)", r.stdout or "", re.M)
    return int(m.group(1)) if m else 0


def _classify_text_only(txt: str) -> str:
    """Статус страницы по одному тексту (без доступа к изображениям).

    Аналог `_classify_page` для poppler. Два отличия:

    1. `blank` не выделяется — страница без текста считается `scan`
       (см. «Граница движка» выше).
    2. Строже порог «нет кириллицы». Базовое правило `_is_garbage` требует
       ≥30 букв, потому что рассчитано на cp1251-кашу, где мусорных букв
       сотни. У poppler отказ выглядит иначе: если в PDF нет ToUnicode-карты
       для подмножества шрифта, глифы **молча выпадают** — остаются цифры и
       пунктуация, букв мало, и до порога 30 дело не доходит. Проверено на
       фикстуре: 293 симв. у pymupdf против 107 у poppler с полностью
       потерянной кириллицей, и страница прошла бы как годный текст.
       Поэтому на fallback-пути: есть буквы, но кириллицы нет вовсе → на
       vision. Ложное срабатывание (целиком латинская страница) стоит одного
       прохода vision — это дешевле потери текста, и та же логика уже принята
       в `_is_garbage`.
    """
    stripped = (txt or "").strip()
    if len(stripped) < 10:
        return "scan"
    diag = _text_diagnostics(txt)
    if _is_garbage(diag):
        return "garbage"
    alpha = diag.get("alpha_count", 0)
    # (а) глифы выпали почти полностью — остались цифры, пробелы и пунктуация.
    #     Ровно этот профиль дала фикстура: 107 симв. текста при ОДНОЙ букве.
    if len(stripped) >= 40 and alpha < 5:
        return "garbage"
    # (б) буквы есть, но кириллицы нет вовсе — частичное выпадение глифов.
    if alpha >= 5 and diag.get("cyr_ratio", 1.0) == 0.0:
        return "garbage"
    return "text"


def _pdf_pages_poppler(filepath: str):
    """(pages, statuses, text_parts) через pdftotext. Бросает исключение при сбое.

    Один вызов на документ: pdftotext разделяет страницы переводом формы
    (\\f), поэтому и число страниц, и постраничный текст берутся из одной
    выдачи. Это заодно снимает зависимость от `pdfinfo` — в некоторых сборках
    poppler (напр. в составе Git for Windows) есть pdftotext, но нет pdfinfo.
    """
    exe = _poppler_bin("pdftotext")
    if not exe:
        raise RuntimeError("poppler (pdftotext) недоступен")

    # -layout сохраняет колоночную вёрстку (важно для таблиц и шапок).
    r = subprocess.run([exe, "-layout", filepath, "-"],
                       capture_output=True, timeout=300)
    if r.returncode != 0:
        err = (r.stderr or b"").decode("utf-8", "replace").strip()
        raise RuntimeError(f"pdftotext rc={r.returncode}: {err[:200]}")

    raw = (r.stdout or b"").decode("utf-8", "replace")
    chunks = raw.split("\f")
    if chunks and not chunks[-1].strip():
        chunks.pop()            # pdftotext ставит \f и после последней страницы
    if not chunks:
        raise RuntimeError("pdftotext вернул пустую выдачу")

    # Сверка с pdfinfo, если он есть: расхождение — сигнал, что разбор по \f
    # разошёлся с реальной пагинацией (не блок, флаг в предупреждениях выше).
    declared = _pdf_page_count_poppler(filepath)
    if declared and declared != len(chunks):
        chunks = chunks[:declared] + [""] * max(0, declared - len(chunks))

    statuses, text_parts = [], []
    for i, txt in enumerate(chunks, start=1):
        status = _classify_text_only(txt)
        statuses.append(status)
        text_parts.append(txt if status == "text" else f"[[VISION_PAGE {i}: {status}]]")
    return len(chunks), statuses, text_parts


def _render_page_poppler(filepath: str, page_index0: int, render_dir: str, width: int) -> dict:
    """Рендер одной страницы в PNG через pdftoppm.

    Плавающий паддинг имени. pdftoppm дописывает номер страницы с ведущими
    нулями по разрядности ОБЩЕГО числа страниц: `p-1.png` для 9-страничного
    документа и `p-01.png` для 10-страничного. Боевой прогон 16.07 на этом
    молча получал 0 символов, потому что искал файл по угаданному имени.
    Поэтому рендерим в отдельный подкаталог и забираем то, что реально
    появилось, а не то, что ожидали.
    """
    exe = _poppler_bin("pdftoppm")
    if not exe:
        return {"path": None, "cropped": False,
                "warnings": ["ни pymupdf, ни poppler (pdftoppm) недоступны"]}

    n = page_index0 + 1
    out_dir = os.path.join(render_dir, f"_p{n}")
    os.makedirs(out_dir, exist_ok=True)
    for stale in glob.glob(os.path.join(out_dir, "*.png")):
        try:
            os.remove(stale)
        except OSError:
            pass

    prefix = os.path.join(out_dir, "page")
    try:
        subprocess.run([exe, "-png", "-scale-to-x", str(width), "-scale-to-y", "-1",
                        "-f", str(n), "-l", str(n), filepath, prefix],
                       capture_output=True, timeout=180, check=True)
    except Exception as e:
        ec, rt, msg = classify_open_error(filepath, e)
        return {"path": None, "cropped": False, "error_class": ec, "retryable": rt,
                "warnings": [f"Рендер через pdftoppm не удался: {msg}"]}

    produced = sorted(glob.glob(os.path.join(out_dir, "page*.png")))
    if not produced:
        return {"path": None, "cropped": False,
                "warnings": [f"pdftoppm не создал PNG для стр. {n}"]}

    # Имя, ожидаемое остальным пайплайном (как у fitz-ветки).
    out = os.path.join(render_dir, f"{Path(filepath).stem}_p{n}.png")
    shutil.move(produced[0], out)
    try:
        os.rmdir(out_dir)
    except OSError:
        pass
    # Guard обрезки (F3.5) здесь недоступен: без геометрии страницы аспект
    # не с чем сравнивать. Помечаем явно, чтобы «не проверено» не читалось
    # как «проверено и ок».
    return {"path": out, "cropped": False, "render_engine": "poppler",
            "warnings": ["рендер через poppler: guard обрезки (F3.5) не применялся"]}


# --- Извлечение из PDF ------------------------------------------------------

def extract_pdf(filepath: str,
                max_head_pages: int = DEFAULT_MAX_HEAD_PAGES,
                structural_threshold: int = DEFAULT_STRUCTURAL_THRESHOLD) -> dict:
    """PDF: программный текст по страницам; решение needs_vision по статусам страниц.

    Vision — основной путь для скан/мусорных страниц (агентская сторона, см. shared/ocr.md).
    Здесь только директива: какие страницы и почему.
    """
    engine_warnings = []
    try:
        import fitz  # pymupdf
        engine = "pymupdf"
    except ImportError:
        # F.2: не отказ, а переход на poppler (предустановлен в песочнице Cowork).
        if not poppler_available():
            return {"text": "", "method": "none", "confidence": "low", "pages": 0,
                    "needs_vision": False, "vision_pages": [], "vision_reason": None,
                    "structural_recommended": False,
                    "warnings": [dependency_hint("pymupdf") + " (либо системный poppler/pdftotext)"]}
        engine = "poppler"
        engine_warnings.append(
            "pymupdf недоступен — извлечение через poppler (pdftotext); "
            "пустые страницы классифицируются как скан и уходят на vision")

    if engine == "poppler":
        try:
            pages, statuses, text_parts = _pdf_pages_poppler(filepath)
        except Exception as e:
            ec, rt, msg = classify_open_error(filepath, e)
            return _pdf_error(ec, rt, msg)
    else:
        try:
            doc = fitz.open(filepath)
        except Exception as e:
            # E.14 а/б: различаем «не материализован (OneDrive) → ретрай» и «реально битый → ветка битого».
            ec, rt, msg = classify_open_error(filepath, e)
            return _pdf_error(ec, rt, msg)

        pages = len(doc)
        statuses = []
        text_parts = []      # текст для 'text'-страниц; плейсхолдер для остальных
        try:
            for i, page in enumerate(doc):
                status = _classify_page(page)
                statuses.append(status)
                if status == "text":
                    text_parts.append(page.get_text())
                elif status == "blank":
                    text_parts.append("")
                else:  # scan / garbage — отдаётся vision
                    text_parts.append(f"[[VISION_PAGE {i + 1}: {status}]]")
        except Exception as e:
            # Файл «открылся, но развалился» при чтении страниц — как правило дегидратация
            # cloud-only на маунте (E.14а): не падаем трейсбеком, а классифицируем.
            doc.close()
            ec, rt, msg = classify_open_error(filepath, e)
            return _pdf_error(ec, rt, msg, pages)
        doc.close()

    # страницы под vision (1-based для человека/агента)
    vision_idx = [i + 1 for i, s in enumerate(statuses) if s in ("scan", "garbage")]
    needs_vision = bool(vision_idx)
    has_garbage = any(s == "garbage" for s in statuses)
    has_scan = any(s == "scan" for s in statuses)

    if needs_vision:
        vision_reason = "garbage-layer" if has_garbage else "no-text-layer"
    else:
        vision_reason = None

    # структурный режим (F3.4): длинный документ под vision → не гнать ВСЕ страницы,
    # а полнотекстить голову + подписную, остальное — структурный скелет (решает агент)
    structural_recommended = needs_vision and pages > structural_threshold
    if structural_recommended:
        head = [i + 1 for i, s in enumerate(statuses)
                if s in ("scan", "garbage") and i < max_head_pages]
        tail = vision_idx[-1] if vision_idx else None
        suggested = sorted(set(head + ([tail] if tail else [])))
    else:
        suggested = vision_idx

    full_text = "\n\n---\n\n".join(text_parts)

    # confidence только для программной части; для vision-страниц confidence ставит агент
    if not needs_vision:
        avg_len = len(full_text) / max(pages, 1)
        confidence = "high" if avg_len > 200 else "medium" if avg_len > 50 else "low"
        method = "pdf-text"
    else:
        # гибрид (часть текстовых, часть vision) или полностью скан
        confidence = "pending-vision"
        method = "pdf-text+vision" if any(s == "text" for s in statuses) else "vision"

    # детект склейки нескольких документов в одном PDF (E.6.2)
    composite = _detect_composite(text_parts, statuses) if pages > 1 else {
        "composite_suspected": False, "composite_reasons": []}

    warnings = list(engine_warnings)   # F.2: чем читали PDF, если не pymupdf
    if has_garbage:
        warnings.append("Обнаружен мусорный текстовый слой — страницы отправлены на vision (F3.3)")
    if structural_recommended:
        warnings.append(
            f"Длинный скан ({pages} стр.) — рекомендован структурный режим (F3.4): "
            f"полнотекст головы (первые {max_head_pages}) + подписной, остальное — скелет")
    if composite["composite_suspected"]:
        warnings.append(
            "Похоже на составной пакет (несколько документов в одном PDF, E.6.2): "
            + "; ".join(composite["composite_reasons"])
            + " — предложить split в preview")

    return {
        "text": full_text,
        "method": method,
        "confidence": confidence,
        "pages": pages,
        "page_statuses": statuses,
        "needs_vision": needs_vision,
        "vision_pages": vision_idx,
        "vision_pages_suggested": suggested,
        "vision_reason": vision_reason,
        "structural_recommended": structural_recommended,
        "composite_suspected": composite["composite_suspected"],
        "composite_reasons": composite["composite_reasons"],
        "warnings": warnings,
    }


# --- Извлечение из DOCX / TXT (без изменений по сути) ------------------------

def extract_docx_text(filepath: str) -> dict:
    """Извлечение текста из DOCX."""
    try:
        from docx import Document
    except ImportError:
        return {"text": "", "method": "none", "confidence": "low", "pages": 0,
                "needs_vision": False, "warnings": [dependency_hint("python-docx")]}

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            text += "\n\n" + "\n".join(rows)

    return {
        "text": text,
        "method": "docx-parse",
        "confidence": "high",
        "pages": max(1, len(paragraphs) // 30),
        "needs_vision": False,
        "warnings": [],
    }


def extract_text_file(filepath: str) -> dict:
    """Чтение текстового файла (UTF-8 → cp1251 fallback)."""
    try:
        with open(filepath, "r", encoding="utf-8-sig") as f:   # -sig срезает BOM, если есть
            text = f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, "r", encoding="cp1251") as f:
                text = f.read()
        except Exception:
            return {"text": "", "method": "none", "confidence": "low", "pages": 0,
                    "needs_vision": False, "warnings": ["Не удалось прочитать файл"]}

    return {
        "text": text,
        "method": "text-read",
        "confidence": "high",
        "pages": 1,
        "needs_vision": False,
        "warnings": [],
    }


# --- Извлечение из таблиц: структурная сводка (G.3) --------------------------
#
# Приём таблицу НЕ анализирует (канон ocr.md: зеркало транскрипционное, юрвыводы
# не в зеркало). Здесь только скелет: листы, шапки, образец строк, итоговые
# строки, диапазон дат, служебные признаки. Арифметика и сверки — analyze_table.py
# на стороне study-evidence.

TABLE_TOTAL_MARKERS = (
    "итого", "всего", "оборот", "сальдо", "баланс", "к оплате", "сумма по",
    "итог", "total",
)
TABLE_MAX_SCAN_ROWS = 200_000     # защита от гигантского листа
TABLE_SAMPLE_ROWS = 15            # строк-образцов в зеркало (по профилю G.3)
TABLE_HEADER_SEARCH_ROWS = 20     # в пределах скольких строк искать шапку
TABLE_FORMULA_SCAN_ROWS = 2000    # предел второго прохода (детект формул без значений)
TABLE_CSV_FULLTEXT_ROWS = 300     # CSV до этого размера отдаём полнотекстом (как раньше)


def _table_error(error_class: str, retryable: bool, message: str) -> dict:
    """Единый JSON-отказ табличного извлечения (контракт E.14)."""
    return {"text": "", "method": "none", "confidence": "low", "pages": 0,
            "needs_vision": False, "error_class": error_class,
            "retryable": retryable, "warnings": [message], "table": None}


def _cell_empty(value) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def _cell_str(value) -> str:
    """Ячейка → строка для зеркала. Даты приводим к ISO, числа не форматируем."""
    import datetime as _dt
    if value is None:
        return ""
    if isinstance(value, _dt.datetime):
        return value.strftime("%Y-%m-%d") if (value.hour, value.minute, value.second) == (0, 0, 0) \
            else value.strftime("%Y-%m-%d %H:%M")
    if isinstance(value, _dt.date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, _dt.time):
        return value.strftime("%H:%M")
    return str(value).strip()


def _is_date_like(value) -> bool:
    import datetime as _dt
    return isinstance(value, (_dt.datetime, _dt.date)) and not isinstance(value, bool)


def _pad_row(cells: list, width: int) -> list:
    """Выравнивает строку до ширины листа. Позиция ячейки = позиция колонки.

    Пустые ячейки НЕ выбрасываются: иначе значения съезжают влево и списание
    печатается на месте прихода (поймано тестом при разработке G.3).
    """
    return (list(cells) + [""] * width)[:width]


def _find_header_row(rows: list) -> int:
    """Индекс строки-шапки (0-based) в пределах первых строк.

    У выписок и реестров сверху 3–7 строк преамбулы (наименование банка, период,
    реквизиты), поэтому первая строка — не шапка. Признак шапки: максимум непустых
    ТЕКСТОВЫХ ячеек при непустой следующей строке. Совпадений нет → 0.
    """
    best_idx, best_score = 0, -1
    limit = min(len(rows), TABLE_HEADER_SEARCH_ROWS)
    for i in range(limit):
        row = rows[i]
        text_cells = sum(1 for c in row if isinstance(c, str) and c.strip())
        if text_cells < 2:
            continue
        next_nonempty = any(not _cell_empty(c) for c in rows[i + 1]) if i + 1 < len(rows) else False
        score = text_cells + (2 if next_nonempty else 0)
        if score > best_score:
            best_idx, best_score = i, score
    return best_idx if best_score > 0 else 0


def _sheet_digest(ws, sheet_name: str, state: str) -> dict:
    """Структурный дайджест одного листа (значения уже вычислены: data_only=True)."""
    rows, truncated = [], False
    for n, row in enumerate(ws.iter_rows(values_only=True)):
        if n >= TABLE_MAX_SCAN_ROWS:
            truncated = True
            break
        rows.append(list(row))

    # хвостовые полностью пустые строки не считаем данными (openpyxl их отдаёт)
    while rows and all(_cell_empty(c) for c in rows[-1]):
        rows.pop()

    if not rows:
        return {"name": sheet_name, "state": state, "rows": 0, "cols": 0,
                "header_row": None, "headers": [], "sample": [], "totals": [],
                "date_range": None, "truncated": truncated}

    cols = max(len(r) for r in rows)
    hdr_idx = _find_header_row(rows)
    headers = _pad_row([_cell_str(c) for c in rows[hdr_idx]], cols)

    # ВАЖНО: пустые ячейки НЕ выбрасываем — иначе значения съезжают влево и, например,
    # списание печатается на месте прихода. Позиция ячейки = позиция колонки, всегда.
    sample = [_pad_row([_cell_str(c) for c in r], cols)
              for r in rows[hdr_idx + 1: hdr_idx + 1 + TABLE_SAMPLE_ROWS]]

    # итоговые строки — по маркеру в любой текстовой ячейке
    totals = []
    for i, row in enumerate(rows):
        if i <= hdr_idx:
            continue
        joined = " ".join(_cell_str(c).lower() for c in row if not _cell_empty(c))
        if joined and any(m in joined for m in TABLE_TOTAL_MARKERS):
            totals.append({"row": i + 1, "cells": _pad_row([_cell_str(c) for c in row], cols)})
        if len(totals) >= 20:
            break

    # диапазон дат по первой колонке, где даты преобладают
    date_range = None
    for col in range(cols):
        values = [r[col] for r in rows[hdr_idx + 1:] if col < len(r)]
        dates = [v for v in values if _is_date_like(v)]
        if dates and len(dates) >= max(3, len(values) // 2):
            date_range = {"column": headers[col] if col < len(headers) else f"col{col + 1}",
                          "min": _cell_str(min(dates)), "max": _cell_str(max(dates))}
            break

    return {"name": sheet_name, "state": state, "rows": len(rows), "cols": cols,
            "header_row": hdr_idx + 1, "headers": headers, "sample": sample,
            "totals": totals, "date_range": date_range, "truncated": truncated}


def _count_stale_formulas(filepath: str, sheet_names: list) -> dict:
    """Формулы без сохранённых значений — сравнение двух проходов по одним координатам.

    Классическая ловушка: файл сохранён без кэша вычисленных значений (типично для
    выгрузок из 1С и банковских систем). openpyxl в режиме значений вернёт пустоту
    там, где в Excel видна сумма, — и суммы молча станут нулями. Поэтому идём по
    файлу дважды параллельно: формулы (data_only=False) против значений
    (data_only=True). Формула есть, значения нет → stale.

    Проверять по «весь лист пуст» недостаточно: пустыми бывают только расчётные
    колонки, а справочные данные при этом читаются — картина выглядит нормальной.
    """
    try:
        from openpyxl import load_workbook
        wb_frm = load_workbook(filepath, data_only=False, read_only=True, keep_links=False)
        wb_val = load_workbook(filepath, data_only=True, read_only=True, keep_links=False)
    except Exception:
        return {"formulas": 0, "stale": 0, "checked": False, "examples": []}

    formulas, stale, examples = 0, 0, []
    try:
        for name in sheet_names:
            if name not in wb_frm.sheetnames or name not in wb_val.sheetnames:
                continue
            frm_rows = wb_frm[name].iter_rows(values_only=True)
            val_rows = wb_val[name].iter_rows(values_only=True)
            for n, (frm_row, val_row) in enumerate(zip(frm_rows, val_rows)):
                if n >= TABLE_FORMULA_SCAN_ROWS:
                    break
                for col, cell in enumerate(frm_row):
                    if not (isinstance(cell, str) and cell.startswith("=")):
                        continue
                    formulas += 1
                    value = val_row[col] if col < len(val_row) else None
                    if _cell_empty(value):
                        stale += 1
                        if len(examples) < 5:
                            examples.append({"sheet": name, "row": n + 1, "col": col + 1,
                                             "formula": cell[:60]})
    except Exception:
        return {"formulas": formulas, "stale": stale, "checked": False, "examples": examples}
    finally:
        for wb in (wb_frm, wb_val):
            try:
                wb.close()
            except Exception:
                pass
    return {"formulas": formulas, "stale": stale, "checked": True, "examples": examples}


TABLE_MD_MAX_COLS = 12   # шире — markdown-таблица нечитаема, переходим на пары «шапка: значение»


def _md_escape(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def _rows_as_md(headers: list, rows: list, row_labels: list = None) -> list:
    """Строки → markdown-таблица (правило 2 mirror-template: «таблицы → markdown-таблицы»).

    Колонки не сдвигаются: пустая ячейка остаётся пустой колонкой.
    """
    out = []
    head = (["стр."] if row_labels else []) + [_md_escape(h) or f"·{i + 1}"
                                               for i, h in enumerate(headers)]
    out.append("| " + " | ".join(head) + " |")
    out.append("|" + "---|" * len(head))
    for i, row in enumerate(rows):
        cells = [_md_escape(c) for c in row]
        if row_labels:
            cells = [str(row_labels[i])] + cells
        out.append("| " + " | ".join(cells) + " |")
    return out


def _rows_as_pairs(headers: list, rows: list, row_labels: list = None) -> list:
    """Широкая таблица → по строке на запись, парами «шапка: значение» (пустые опускаем)."""
    out = []
    for i, row in enumerate(rows):
        label = f"строка {row_labels[i]}" if row_labels else f"запись {i + 1}"
        pairs = [f"{headers[j] or f'кол.{j + 1}'}: {c}"
                 for j, c in enumerate(row) if c]
        out.append(f"- **{label}** — " + "; ".join(pairs))
    return out


def _table_summary_text(table: dict, filename: str, warnings: list = None) -> str:
    """Человекочитаемая структурная сводка — тело зеркала.

    Критичные предупреждения печатаются **в теле**, а не только в поле `warnings`
    выхода: ни один скилл-приёмник поле `warnings` не читает (проверено разведкой
    2026-07-30), поэтому предупреждение, оставленное только там, не дойдёт ни до
    юриста, ни до нижестоящих скиллов.
    """
    lines = [f"# Структурная сводка таблицы: {filename}", ""]

    if warnings:
        lines.append("## ⚠️ Требует внимания")
        lines.append("")
        for w in warnings:
            lines.append(f"- {w}")
        lines.append("")

    lines.append(f"Листов: {len(table['sheets'])}"
                 + (f" (из них скрытых: {table['hidden_sheets']})" if table["hidden_sheets"] else ""))
    if table.get("formulas"):
        lines.append(f"Ячеек с формулами (в пределах проверки): {table['formulas']}"
                     + (f", из них без сохранённого значения: {table['stale_formulas']}"
                        if table.get("stale_formulas") else ""))
    lines.append("")
    lines.append("> Значения в сводке прочитаны машинно и точны — но тело таблицы **неполно**: "
                 "вынесены шапка, первые строки и итоговые строки. Суммы по всему массиву, "
                 "сверки и поиск расхождений — `study-evidence` (`scripts/analyze_table.py`).")
    lines.append("")

    for sh in table["sheets"]:
        title = f"## Лист «{sh['name']}»"
        if sh["state"] != "visible":
            title += f"  ⚠️ СКРЫТЫЙ ({sh['state']})"
        lines.append(title)
        lines.append(f"Строк: {sh['rows']}, колонок: {sh['cols']}"
                     + (" — **лист усечён при чтении**" if sh.get("truncated") else ""))
        if sh["rows"] == 0:
            lines += ["", "_Лист пуст._", ""]
            continue
        if sh["date_range"]:
            dr = sh["date_range"]
            lines.append(f"Диапазон дат (колонка «{dr['column']}»): {dr['min']} — {dr['max']}")
        lines.append("")

        narrow = sh["cols"] <= TABLE_MD_MAX_COLS
        headers = sh["headers"]

        if sh["sample"]:
            lines.append(f"**Шапка — строка {sh['header_row']}; далее первые "
                         f"{len(sh['sample'])} строк(и):**")
            lines.append("")
            first_data_row = sh["header_row"] + 1
            labels = list(range(first_data_row, first_data_row + len(sh["sample"])))
            lines += (_rows_as_md(headers, sh["sample"], labels) if narrow
                      else _rows_as_pairs(headers, sh["sample"], labels))
            lines.append("")
        elif headers:
            lines.append(f"**Шапка (строка {sh['header_row']}):** "
                         + " | ".join(h for h in headers if h))
            lines.append("")

        if sh["totals"]:
            lines.append("**Итоговые строки, найденные по маркеру:**")
            lines.append("")
            labels = [t["row"] for t in sh["totals"]]
            rows = [t["cells"] for t in sh["totals"]]
            lines += (_rows_as_md(headers, rows, labels) if narrow
                      else _rows_as_pairs(headers, rows, labels))
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _user_site_dirs() -> list:
    """Каталоги пользовательских пакетов (`AppData\\Roaming\\Python\\...` на Windows)."""
    dirs = []
    try:
        import site
        if hasattr(site, "getusersitepackages"):
            value = site.getusersitepackages()
            dirs += [value] if isinstance(value, str) else list(value)
    except Exception:
        pass
    return [d for d in dirs if d]


def ensure_user_site() -> list:
    """Подключить user-site, если он существует, но исключён из `sys.path`.

    Боевой случай (2026-07-30, машина Сюзерена): `openpyxl` физически установлен в
    `C:\\Users\\…\\AppData\\Roaming\\Python\\Python314\\site-packages`, но в его сессии
    PowerShell этот каталог в `sys.path` не попадает — и скрипт отвечает «не
    установлен» при установленной библиотеке. Причины бывают разные
    (`PYTHONNOUSERSITE`, запуск с `-s`, политика окружения), и выяснять их на каждой
    машине бессмысленно: если каталог есть на диске, подключаем его сами.
    """
    added = []
    for d in _user_site_dirs():
        if os.path.isdir(d) and d not in sys.path:
            sys.path.append(d)
            added.append(d)
    return added


ensure_user_site()      # до первых импортов сторонних библиотек


def _module_on_disk(module: str) -> str:
    """Путь к пакету в user-site, если он там физически лежит (иначе '')."""
    name = module.replace("-", "_").lower()
    aliases = {"python_docx": "docx", "pymupdf": "fitz"}
    name = aliases.get(name, name)
    for d in _user_site_dirs():
        for candidate in (os.path.join(d, name),
                          os.path.join(d, name + ".py")):
            if os.path.exists(candidate):
                return candidate
    return ""


def dependency_hint(module: str) -> str:
    """Сообщение о зависимости — с интерпретатором и готовой командой установки.

    «Запустите setup.sh» бесполезно там, где проблема и возникает: на Windows в
    PowerShell bash-скрипт не исполняется, а `python3` разрешается в ДРУГОЙ
    интерпретатор (заглушка WindowsApps) без зависимостей — при том что в `python`
    они стоят. Молчаливый уход в fallback при установленной зависимости — тот же
    класс дефекта, что F.26 (тихая деградация при верной настройке).
    """
    on_disk = _module_on_disk(module)
    if on_disk:
        # Пакет есть, но не подхватился даже после ensure_user_site() — значит каталог
        # исключён жёстко. Тогда «установите» — вредный совет: он уводит от причины.
        return (f"{module} УСТАНОВЛЕН, но недоступен для импорта.\n"
                f"  Пакет лежит: {on_disk}\n"
                f"  Запущено:    {sys.executable}\n"
                f"  Каталог пользовательских пакетов исключён из пути поиска. "
                f"Проверьте: \"{sys.executable}\" -c \"import site,sys; "
                f"print(site.ENABLE_USER_SITE, sys.flags.no_user_site)\" — если там "
                f"False/1, снимите PYTHONNOUSERSITE (`$env:PYTHONNOUSERSITE=''`) либо "
                f"установите пакет в системный каталог: "
                f"\"{sys.executable}\" -m pip install --target "
                f"\"{os.path.join(sys.prefix, 'Lib', 'site-packages')}\" {module}")
    return (f"{module} не установлен для этого интерпретатора.\n"
            f"  Запущено: {sys.executable}\n"
            f"  Установить: \"{sys.executable}\" -m pip install {module}\n"
            f"  Либо запустите тем интерпретатором, где зависимости есть (на Windows "
            f"обычно `python`, а не `python3`); в Linux/Cowork — scripts/setup.sh")


def extract_xlsx_structure(filepath: str) -> dict:
    """.xlsx / .xlsm → структурная сводка (не анализ). Контракт как у остальных."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return _table_error("corrupt", False, dependency_hint("openpyxl"))

    try:
        wb = load_workbook(filepath, data_only=True, read_only=True, keep_links=False)
    except Exception as exc:
        error_class, retryable, message = classify_open_error(filepath, exc)
        # Порядок важен: 0-байтовый и не материализованный файл имеют свои ветки приёма
        # (E.14 / E.6) и «паролем/битым» называться не должны — иначе пустышку отправят
        # снимать несуществующий пароль вместо запроса перезалива.
        if error_class in ("empty", "not_materialized"):
            return _table_error(error_class, retryable, message.replace("PDF", "Файл"))
        msg = str(exc).lower()
        # зашифрованный .xlsx и переименованный .xls — оба не zip-контейнеры
        if "not a zip" in msg or "badzipfile" in type(exc).__name__.lower():
            return _table_error(
                "corrupt", False,
                "Файл не является .xlsx: возможно, он защищён паролем либо это старый "
                "формат .xls с новым расширением. Снимите пароль или сохраните как .xlsx")
        return _table_error(error_class, retryable, message.replace("PDF", "Файл"))

    try:
        sheets, hidden = [], 0
        for name in wb.sheetnames:
            ws = wb[name]
            state = getattr(ws, "sheet_state", "visible") or "visible"
            if state != "visible":
                hidden += 1
            sheets.append(_sheet_digest(ws, name, state))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    formula_info = _count_stale_formulas(filepath, [s["name"] for s in sheets])

    table = {
        "kind": "spreadsheet",
        "sheets": sheets,
        "hidden_sheets": hidden,
        "formulas": formula_info["formulas"],
        "stale_formulas": formula_info["stale"],
        "stale_examples": formula_info["examples"],
        "total_rows": sum(s["rows"] for s in sheets),
    }

    warnings = []
    # Ловушка «формулы без сохранённых значений» — первой строкой: она обесценивает суммы.
    if formula_info["stale"]:
        where = ", ".join(f"«{e['sheet']}» стр. {e['row']}" for e in formula_info["examples"][:3])
        warnings.append(
            f"**Формулы без сохранённых значений: {formula_info['stale']} ячеек** ({where}). "
            f"В Excel вы видите там суммы, а в файле их нет — программное чтение вернёт пустоту, "
            f"и итоги окажутся заниженными. Откройте файл в Excel, пересохраните и повторите приём. "
            f"Типично для выгрузок из 1С и банковских систем")
    if hidden:
        warnings.append(f"Скрытых листов: {hidden} — проверьте их содержимое: там бывает "
                        f"исходник до правки")
    if any(s.get("truncated") for s in sheets):
        warnings.append(f"Лист(ы) усечены при чтении (предел {TABLE_MAX_SCAN_ROWS} строк) — "
                        f"итоги по всему массиву считать через `analyze_table.py`")
    if table["total_rows"] == 0:
        warnings.append("Все листы пусты — проверьте файл (возможно, это шаблон или выгрузка "
                        "не сохранилась)")
    if not formula_info["checked"] and formula_info["formulas"] == 0:
        warnings.append("Проверка формул не выполнена (файл не открылся вторым проходом) — "
                        "если в Excel видны расчётные колонки, сверьте итоги вручную")

    return {
        "text": _table_summary_text(table, Path(filepath).name, warnings),
        "method": "table-structure",
        "confidence": "high",
        "pages": max(1, len(sheets)),
        "needs_vision": False,
        "warnings": warnings,
        "table": table,
    }


def extract_csv_table(filepath: str) -> dict:
    """CSV: небольшой — полнотекстом (как раньше), крупный — структурной сводкой."""
    import csv as _csv

    raw = None
    for encoding in ("utf-8-sig", "cp1251"):
        try:
            with open(filepath, "r", encoding=encoding, newline="") as f:
                raw = f.read()
            break
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            error_class, retryable, message = classify_open_error(filepath, exc)
            return _table_error(error_class, retryable, message.replace("PDF", "Файл"))
    if raw is None:
        return _table_error("corrupt", False, "Не удалось прочитать файл ни в UTF-8, ни в cp1251")

    try:
        dialect = _csv.Sniffer().sniff(raw[:4096], delimiters=";,\t|")
        delimiter = dialect.delimiter
    except Exception:
        delimiter = ";" if raw[:4096].count(";") > raw[:4096].count(",") else ","

    rows = [r for r in _csv.reader(raw.splitlines(), delimiter=delimiter)]
    while rows and all(not c.strip() for c in rows[-1]):
        rows.pop()

    # Мелкий CSV дешевле и полезнее отдать целиком — поведение как до G.3.
    if len(rows) <= TABLE_CSV_FULLTEXT_ROWS:
        return {"text": raw, "method": "text-read", "confidence": "high", "pages": 1,
                "needs_vision": False,
                "warnings": [f"CSV прочитан полнотекстом ({len(rows)} строк, разделитель "
                             f"«{delimiter}»)"],
                "table": {"kind": "csv", "rows": len(rows), "delimiter": delimiter,
                          "fulltext": True}}

    hdr_idx = _find_header_row(rows)
    sheet = {
        "name": Path(filepath).name, "state": "visible", "rows": len(rows),
        "cols": max(len(r) for r in rows), "header_row": hdr_idx + 1,
        "headers": [c.strip() for c in rows[hdr_idx]],
        "sample": [[c.strip() for c in r]
                   for r in rows[hdr_idx + 1: hdr_idx + 1 + TABLE_SAMPLE_ROWS]],
        "totals": [{"row": i + 1, "cells": [c.strip() for c in r if c.strip()]}
                   for i, r in enumerate(rows)
                   if i > hdr_idx
                   and any(m in " ".join(r).lower() for m in TABLE_TOTAL_MARKERS)][:20],
        "date_range": None, "truncated": False,
    }
    table = {"kind": "csv", "sheets": [sheet], "hidden_sheets": 0, "formulas": 0,
             "delimiter": delimiter, "total_rows": len(rows), "fulltext": False}
    csv_warn = [f"CSV крупный ({len(rows)} строк) — в зеркало вынесена структурная сводка, "
                f"не полный текст; итоги по массиву — через `analyze_table.py`"]
    return {
        "text": _table_summary_text(table, Path(filepath).name, csv_warn),
        "method": "table-structure",
        "confidence": "high",
        "pages": 1,
        "needs_vision": False,
        "warnings": [f"CSV крупный ({len(rows)} строк) — в зеркало вынесена структурная "
                     f"сводка, не полный текст"],
        "table": table,
    }


def extract_xls_legacy(filepath: str) -> dict:
    """.xls (формат до 2007) — программно не читаем: честный отказ, не «неизвестный формат»."""
    return {"text": "", "method": "none", "confidence": "low", "pages": 0,
            "needs_vision": False, "table": None,
            "warnings": ["Формат .xls (Excel до 2007) не поддерживается: откройте файл и "
                         "сохраните как .xlsx, затем повторите приём. Библиотеки xlrd в "
                         "среде нет by-design — новый формат покрывает все боевые случаи"]}


def extract_image(filepath: str) -> dict:
    """Изображение: vision — основной путь (не tesseract). Директива агенту."""
    return {
        "text": "",
        "method": "vision",
        "confidence": "pending-vision",
        "pages": 1,
        "needs_vision": True,
        "vision_pages": [1],
        "vision_pages_suggested": [1],
        "vision_reason": "image",
        "structural_recommended": False,
        "warnings": ["Изображение → vision (Read tool), см. shared/ocr.md"],
    }


# --- Рендер страницы в PNG для vision (F3.1) + guard обрезки (F3.5) ----------

def render_page_to_png(filepath: str, page_index0: int, render_dir: str,
                       width: int = DEFAULT_RENDER_WIDTH) -> dict:
    """Рендер одной страницы PDF в PNG через pymupdf.

    pymupdf рендерит детерминированно в процессе (без системного pdftoppm,
    который у файлового Read падает «unsafe location», F2.11). page_index0 — 0-based.
    Возвращает путь PNG + флаг cropped (сверка аспекта рендера и страницы, F3.5).
    """
    try:
        import fitz
    except ImportError:
        # F.2: рендер — вторая нога, которая раньше отваливалась вместе с pymupdf
        # и уносила с собой vision (ему нечего было читать).
        return _render_page_poppler(filepath, page_index0, render_dir, width)

    try:
        doc = fitz.open(filepath)
        if page_index0 < 0 or page_index0 >= len(doc):
            doc.close()
            return {"path": None, "cropped": False,
                    "warnings": [f"Страница {page_index0 + 1} вне диапазона (всего {len(doc)})"]}
        page = doc[page_index0]
        rect = page.rect
        zoom = (width / rect.width) if rect.width else 2.0
        zoom = max(0.5, min(zoom, 4.0))    # разумные границы downscale/upscale
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        os.makedirs(render_dir, exist_ok=True)
        out = os.path.join(render_dir, f"{Path(filepath).stem}_p{page_index0 + 1}.png")
        pix.save(out)
        # guard обрезки: аспект PNG должен совпасть с аспектом страницы
        page_aspect = (rect.width / rect.height) if rect.height else 0
        img_aspect = (pix.width / pix.height) if pix.height else 0
        cropped = page_aspect > 0 and abs(page_aspect - img_aspect) > 0.02 * page_aspect
        doc.close()
        warnings = []
        if cropped:
            warnings.append(
                f"cropped_render: PNG стр. {page_index0 + 1} обрезан "
                f"(аспект {img_aspect:.3f} ≠ страница {page_aspect:.3f}) — проверить рендер (F3.5)")
        return {"path": out, "width": pix.width, "height": pix.height,
                "cropped": cropped, "warnings": warnings}
    except Exception as e:
        # E.14а: тот же разбор — не материализован (ретрай) vs битый.
        ec, rt, msg = classify_open_error(filepath, e)
        return {"path": None, "cropped": False, "error_class": ec, "retryable": rt,
                "warnings": [f"Рендер не удался: {msg}"]}


# --- Tesseract спот-сверка критичного поля (F3.1; оживлена в F.18) -----------
#
# До F.18 функция была достижима только из shell вручную: вход — готовый КРОП поля,
# которого вызывающему негде было взять (геометрию поля vision не отдаёт), а сравнение
# «прочитано агентом vs прочитано движком» оставалось на глаз той же модели, что и
# транскрибировала. Поэтому здесь появились две вещи: (1) вход «PDF + номер страницы»
# (кроп — необязательное сужение, в долях страницы), (2) машинный троичный вердикт
# по --expect. Смысл ноги — второе чтение с ДРУГИМИ режимами отказа, чем у vision.

_SPOT_TIMEOUT_CROP = 30              # кроп одной строки
_SPOT_TIMEOUT_PAGE = 120             # целая страница при ~1500 px, с запасом
_SPACE_CHARS = "     ⁠"
_NUM_TOKEN_RE = re.compile(r"\d[\d" + _SPACE_CHARS + r"\s.,]*\d|\d")


def find_tessdata() -> str:
    """Путь к вендоренным tessdata: $VASSAL_TESSDATA или scripts/tessdata/."""
    env = os.environ.get("VASSAL_TESSDATA")
    if env and os.path.exists(os.path.join(env, "rus.traineddata")):
        return env
    vendored = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tessdata")
    if os.path.exists(os.path.join(vendored, "rus.traineddata")):
        return vendored
    return ""


def _num_key(value) -> str:
    """Ключ сравнения числового реквизита: цифры + значащая дробная часть.

    «3 390 000,00» → «3390000»; «12.03.2026» → «12032026»; «15,5» → «15.5».
    Ведущие нули НЕ срезаются (ИНН регионов 01–09 начинается с нуля). Нечисловое → "".
    """
    t = re.sub(r"[" + _SPACE_CHARS + r"\s]+", "", str(value or ""))
    if not t or not re.fullmatch(r"[\d.,]+", t):
        return ""
    m = re.fullmatch(r"(.*?)([.,])(\d{1,2})$", t)     # последний разделитель = десятичный
    intpart, frac = (m.group(1), m.group(3)) if m else (t, "")
    intpart = re.sub(r"[.,]", "", intpart)            # остальные — разряды тысяч
    if not intpart.isdigit():
        return ""
    frac = frac.rstrip("0")
    return intpart + ("." + frac if frac else "")


def _text_key(value) -> str:
    """Ключ сравнения нечислового реквизита: нижний регистр, ё→е, только буквы/цифры."""
    return re.sub(r"[^0-9a-zа-я]+", "", str(value or "").lower().replace("ё", "е"))


def _spot_is_blind(ocr_text: str) -> bool:
    """Движок сам ничего внятного не прочитал (пусто/каша) → вердикт inconclusive.

    Считаем содержательными буквы и цифры: на кропе одного реквизита букв может не
    быть вовсе («3 390 000,00»), и объявлять такое чтение слепым нельзя — иначе
    расхождение именно там, где сверка точнее всего, превратится в inconclusive.
    Порог `_is_garbage` рассчитан на текстовый слой в сотни символов, поэтому один
    он тут не работает: страница, с которой tesseract снял только пунктуацию,
    его не срабатывает и проходила бы как «прочитанная».
    """
    t = (ocr_text or "").strip()
    letters = sum(1 for ch in t if ch.isalpha())
    digits = sum(1 for ch in t if ch.isdigit())
    if letters + digits < 2:
        return True
    return _is_garbage(_text_diagnostics(t))


def compare_field(expected, ocr_text: str) -> dict:
    """Машинная сверка «прочитано агентом vs прочитано tesseract» (F.18).

    Сверяет ПРОГРАММА, а не модель — иначе петля самооценки просто удлиняется.
    Вердикт троичный: `inconclusive` (движок ослеп на этой странице) — НЕ повод
    поднимать расхождение, иначе слепой tesseract зашумит флагами каждый документ.
    """
    numeric_key = _num_key(expected)
    engine_numbers = []
    if numeric_key:
        found = {k for k in (_num_key(tok) for tok in _NUM_TOKEN_RE.findall(ocr_text or "")) if k}
        matched = numeric_key in found
        engine_numbers = sorted(found)[:12]
    else:
        key = _text_key(expected)
        matched = bool(key) and key in _text_key(ocr_text)
    if matched:
        verdict = "match"
    elif _spot_is_blind(ocr_text):
        verdict = "inconclusive"
    else:
        verdict = "mismatch"
    return {"expected": str(expected), "verdict": verdict,
            "numeric": bool(numeric_key), "engine_numbers": engine_numbers}


def _render_for_spot(pdf_path: str, page_index0: int, render_dir: str, region=None) -> dict:
    """PNG страницы (или её области) под спот-сверку. region — доли страницы 0..1."""
    if not region:
        res = render_page_to_png(pdf_path, page_index0, render_dir, DEFAULT_RENDER_WIDTH)
        res["regioned"] = False
        return res
    try:
        import fitz
    except ImportError:
        res = render_page_to_png(pdf_path, page_index0, render_dir, DEFAULT_RENDER_WIDTH)
        res.setdefault("warnings", []).append(
            "region требует pymupdf — отрендерена страница целиком (poppler кроп не умеет)")
        res["regioned"] = False
        return res
    try:
        doc = fitz.open(pdf_path)
        if page_index0 < 0 or page_index0 >= len(doc):
            n = len(doc)
            doc.close()
            return {"path": None, "regioned": False,
                    "warnings": [f"Страница {page_index0 + 1} вне диапазона (всего {n})"]}
        pg = doc[page_index0]
        r = pg.rect
        x0, y0, x1, y1 = region
        clip = fitz.Rect(r.x0 + x0 * r.width, r.y0 + y0 * r.height,
                         r.x0 + x1 * r.width, r.y0 + y1 * r.height)
        # кроп мельче страницы — поднимаем масштаб, иначе мелкая цифра не читается
        zoom = max(1.0, min(DEFAULT_RENDER_WIDTH / max(clip.width, 1.0), 8.0))
        pix = pg.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=clip)
        os.makedirs(render_dir, exist_ok=True)
        out = os.path.join(render_dir, f"{Path(pdf_path).stem}_p{page_index0 + 1}_spot.png")
        pix.save(out)
        doc.close()
        return {"path": out, "regioned": True, "width": pix.width, "height": pix.height,
                "warnings": []}
    except Exception as e:
        ec, rt, msg = classify_open_error(pdf_path, e)
        return {"path": None, "regioned": False, "error_class": ec, "retryable": rt,
                "warnings": [f"Рендер области не удался: {msg}"]}


def spot_check_field(target: str, page: int = 1, region=None, expect="",
                     psm: int = 0, tessdata_dir: str = "", render_dir: str = "") -> dict:
    """Независимое второе чтение критичного поля движком tesseract (не нейросетью).

    Вход — PDF + номер страницы (1-based; рендерим сами) ИЛИ готовый PNG/JPG.
    `region` (x0,y0,x1,y1 в долях страницы) сужает область — запасной путь, когда
    целая страница упирается в таймаут; кроп работает только на pymupdf.
    `expect` — то, что прочитал агент: сверка машинная (см. compare_field).

    Не бросает исключений и НЕ блокирует приём: нет словаря / бинаря / таймаут →
    available:false, вызывающий продолжает с результатом vision (shared/ocr.md §9).
    """
    warnings = []
    tessdata_dir = tessdata_dir or find_tessdata()
    if not tessdata_dir:
        return {"available": False, "text": "", "verdict": None,
                "warnings": ["rus.traineddata не найден (scripts/tessdata/ или $VASSAL_TESSDATA) — спот-сверка пропущена"]}

    image_path, full_page = target, True
    if Path(target).suffix.lower() == ".pdf":
        rd = render_dir or os.path.join(os.getcwd(), "outputs", "spot")
        rendered = _render_for_spot(target, max(1, page) - 1, rd, region)
        warnings.extend(rendered.get("warnings", []))
        if not rendered.get("path"):
            warnings.append("рендер страницы не удался — спот-сверка пропущена")
            return {"available": False, "text": "", "verdict": None, "warnings": warnings}
        image_path = rendered["path"]
        full_page = not rendered.get("regioned")
    elif region:
        warnings.append("region задан для готового изображения — игнорируется (кроп делается при рендере из PDF)")

    psm = psm or (3 if full_page else 7)              # 3 — авторазметка страницы, 7 — одна строка
    timeout = _SPOT_TIMEOUT_PAGE if full_page else _SPOT_TIMEOUT_CROP
    env = dict(os.environ)
    env["TESSDATA_PREFIX"] = tessdata_dir             # прямой вызов словарь не увидит (F.5)
    try:
        result = subprocess.run(
            ["tesseract", image_path, "stdout", "-l", "rus", "--psm", str(psm)],
            capture_output=True, text=True, timeout=timeout, env=env
        )
    except FileNotFoundError:
        warnings.append("бинарь tesseract не найден — спот-сверка пропущена")
        return {"available": False, "text": "", "verdict": None, "warnings": warnings}
    except subprocess.TimeoutExpired:
        warnings.append(f"tesseract таймаут {timeout} c — спот-сверка пропущена; "
                        f"сузьте область через --region или продолжайте без неё")
        return {"available": False, "text": "", "verdict": None, "warnings": warnings}

    text = (result.stdout or "").strip()
    out = {"available": True, "text": text, "psm": psm, "image": image_path,
           "verdict": None, "warnings": warnings}
    if str(expect or "").strip():
        out.update(compare_field(expect, text))
    return out


# --- Диспетчер --------------------------------------------------------------

def extract(filepath: str,
            max_head_pages: int = DEFAULT_MAX_HEAD_PAGES,
            structural_threshold: int = DEFAULT_STRUCTURAL_THRESHOLD) -> dict:
    """Определяет тип файла и вызывает нужный экстрактор. Добавляет content_hash."""
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        result = extract_pdf(filepath, max_head_pages, structural_threshold)
    elif ext == ".docx":
        result = extract_docx_text(filepath)
    elif ext in (".xlsx", ".xlsm"):
        result = extract_xlsx_structure(filepath)          # G.3
    elif ext == ".csv":
        result = extract_csv_table(filepath)               # G.3 (мелкий CSV — как раньше)
    elif ext == ".xls":
        result = extract_xls_legacy(filepath)              # G.3
    elif ext in (".txt", ".md", ".html", ".htm", ".xml", ".json", ".yaml", ".yml"):
        result = extract_text_file(filepath)
    elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"):
        result = extract_image(filepath)
    elif ext in (".doc", ".rtf", ".odt"):
        result = {"text": "", "method": "none", "confidence": "low", "pages": 0,
                  "needs_vision": False,
                  "warnings": [f"Формат {ext} не поддерживается программно. Используй Read tool."]}
    else:
        result = {"text": "", "method": "none", "confidence": "low", "pages": 0,
                  "needs_vision": False, "warnings": [f"Неизвестный формат: {ext}"]}

    result["content_hash"] = compute_content_hash(filepath)
    # нормализуем поля контракта (для единообразия выхода)
    result.setdefault("needs_vision", False)
    result.setdefault("vision_pages", [])
    result.setdefault("vision_pages_suggested", result.get("vision_pages", []))
    result.setdefault("vision_reason", None)
    result.setdefault("structural_recommended", False)
    result.setdefault("composite_suspected", False)   # E.6.2 — только PDF взводит
    result.setdefault("composite_reasons", [])
    result.setdefault("low_confidence_fields", [])   # заполняется vision-ногой агента
    result.setdefault("error_class", None)           # E.14 — not_materialized|corrupt|empty (иначе None)
    result.setdefault("retryable", False)            # E.14 — True → ретрай с паузой (не «битый»)
    result.setdefault("table", None)                 # G.3 — структура таблицы (не PDF/DOCX → None)
    return result


# --- CLI --------------------------------------------------------------------

def _get_opt(argv, name, default=None):
    if name in argv:
        i = argv.index(name)
        if i + 1 < len(argv):
            return argv[i + 1]
    return default


def _parse_region(raw: str):
    """«x0,y0,x1,y1» в долях страницы (0..1) → кортеж. Мусор → ValueError."""
    vals = [float(p) for p in str(raw).replace(" ", "").split(",")]
    if len(vals) != 4:
        raise ValueError("нужно ровно 4 числа")
    x0, y0, x1, y1 = vals
    if not (0 <= x0 < x1 <= 1 and 0 <= y0 < y1 <= 1):
        raise ValueError("координаты вне 0..1 или x0≥x1 / y0≥y1")
    return (x0, y0, x1, y1)


def main():
    # Печатаем UTF-8 независимо от консоли ОС (Windows cp1251 иначе рвёт кириллицу/BOM;
    # в Cowork/Linux безвредно — там stdout уже UTF-8).
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

    argv = sys.argv[1:]
    if not argv:
        print("Использование: python3 extract_text.py <файл> [--output-dir DIR] [--render-dir DIR]")
        print("               python3 extract_text.py --render <pdf> --page N [--render-dir DIR] [--width W]")
        print("               python3 extract_text.py --spot-check <pdf|png> [--page N] [--expect ЗНАЧЕНИЕ]")
        print("                                       [--region x0,y0,x1,y1] [--psm N] [--tessdata DIR]")
        sys.exit(1)

    # Режим 3: спот-сверка критичного поля вторым движком (F.18)
    if "--spot-check" in argv:
        target = _get_opt(argv, "--spot-check")
        if not target or not os.path.exists(target):
            print(json.dumps({"error": f"Файл не найден: {target}"}, ensure_ascii=False))
            sys.exit(1)
        region = None
        if "--region" in argv:
            try:
                region = _parse_region(_get_opt(argv, "--region", ""))
            except (ValueError, AttributeError) as e:
                print(json.dumps({"error": f"--region: {e}; формат «x0,y0,x1,y1» в долях страницы"},
                                 ensure_ascii=False))
                sys.exit(1)
        print(json.dumps(spot_check_field(
            target,
            page=int(_get_opt(argv, "--page", "1")),
            region=region,
            expect=_get_opt(argv, "--expect", "") or "",
            psm=int(_get_opt(argv, "--psm", "0")),
            tessdata_dir=_get_opt(argv, "--tessdata", "") or "",
            render_dir=_get_opt(argv, "--render-dir", "") or "",
        ), ensure_ascii=False, indent=2))
        return

    # Режим 2: рендер одной страницы
    if "--render" in argv:
        pdf = _get_opt(argv, "--render")
        if not pdf or not os.path.exists(pdf):
            print(json.dumps({"error": f"Файл не найден: {pdf}"}, ensure_ascii=False))
            sys.exit(1)
        page = int(_get_opt(argv, "--page", "1"))
        render_dir = _get_opt(argv, "--render-dir", os.path.join(os.getcwd(), "outputs", "renders"))
        width = int(_get_opt(argv, "--width", str(DEFAULT_RENDER_WIDTH)))
        print(json.dumps(render_page_to_png(pdf, page - 1, render_dir, width),
                         ensure_ascii=False, indent=2))
        return

    # Режим 1: извлечение
    filepath = argv[0]
    if not os.path.exists(filepath):
        # E.14а: на облачном маунте отсутствие enumerated-файла — почти всегда дегидратация
        # (скилл приёма передаёт реально перечисленные файлы), а не «нет такого файла».
        # Отдаём структурный JSON с retryable, чтобы скилл сделал ретрай с паузой, а не «битый».
        print(json.dumps({
            "text": "", "method": "none", "confidence": "low", "pages": 0,
            "needs_vision": False, "content_hash": "",
            "error_class": "not_materialized", "retryable": True,
            "warnings": [f"Файл не найден на маунте: {filepath} — возможно, не синхронизирован "
                         f"OneDrive (cloud-only); сначала Read (материализует), затем ретрай с паузой"],
        }, ensure_ascii=False, indent=2))
        return

    max_head = int(_get_opt(argv, "--max-head-pages", str(DEFAULT_MAX_HEAD_PAGES)))
    threshold = int(_get_opt(argv, "--structural-threshold", str(DEFAULT_STRUCTURAL_THRESHOLD)))
    result = extract(filepath, max_head, threshold)

    output_dir = _get_opt(argv, "--output-dir")
    if output_dir and result.get("text"):
        os.makedirs(output_dir, exist_ok=True)
        txt_path = os.path.join(output_dir, Path(filepath).stem + ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(result["text"])
        result["saved_to"] = txt_path

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
