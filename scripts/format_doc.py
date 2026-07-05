#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
format_doc.py — детерминированная нога рендера процессуального документа .md -> .docx.

Оформляет готовый подаваемый текст по фирменному стандарту типографики
(см. skills/format-doc/references/style-spec.md). Headless, без Word-аддина,
на python-docx. Тот же вход -> тот же выход (идемпотентно: метаданные фиксированы).

Использование:
    python3 format_doc.py <in.md> <out.docx> --type <тип> [--case case.yaml]

Типы (--type):
    appeal | cassation | nadzor            -> «жалоба»   (ПРОШУ вверху, мини-разделы)
    otzyv | hodataystvo | zayavlenie |
    vozrazhenie | poyasneniya              -> «иск-документ» (ПРОШУ внизу, один список)
    pretenziya                             -> «претензия» (адресат — контрагент)
Русские названия («жалоба», «отзыв», «претензия» …) тоже принимаются.

⚠️ СТАТУС: v1, требует валидации на реальном .md в Cowork (python-docx локально
недоступен). Узлы наибольшего риска (см. style-spec §8): OOXML-нумерация с рестартом
по блокам, границы абзацев (Heading2/цитата), заливки ячеек, character spacing в шапке,
highlight #D9D9D9. При отладке проверять именно их.
"""

import sys
import re
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "format_doc: python-docx недоступен. Оформление .docx пропущено "
        "(fallback: остаётся .md). Установите: pip install python-docx.\n"
    )
    sys.exit(3)  # код 3 -> скилл трактует как graceful fallback


# --------------------------------------------------------------------------- #
# Константы стиля (из style-spec.md §1)
# --------------------------------------------------------------------------- #
FONT = "Garamond"
TABLE_FONT = "Calibri"
BLACK = RGBColor(0x00, 0x00, 0x00)
HL_FILL = "D9D9D9"          # highlight ключевых фраз + заголовок таблицы
TOTAL_FILL = "F2F2F2"       # итоговая строка таблицы

SZ_BODY = 11
SZ_H1 = 16
SZ_H2 = 14
SZ_MINIHEAD = 12
SZ_SUBHEAD = 11
SZ_PROSBA = 14
SZ_TABLE = 9

# межстрочный: style-spec «13.8pt (множитель 1.15)». Берём множитель 1.15 (WD MULTIPLE);
# если в Cowork визуально не совпадёт с эталоном proc-doc-style — переключить на
# EXACTLY Pt(13.8) (см. _para).
LINE_MULT = 1.15

HEADER_LEFT_INDENT = Pt(210)   # шапка прижата к правому блоку
QUOTE_INDENT = Pt(35.4)
LIST_LEFT = Pt(36)
LIST_HANG = Pt(18)

# Маркеры блоков в .md (якорный контракт, style-spec §2 / SKILL.md)
RE_DATE = re.compile(r"^\s*(\d{2}\.\d{2}\.\d{4})\s*$")
RE_H1 = re.compile(r"^\s*#\s+(.*\S)\s*$")
RE_H2 = re.compile(r"^\s*##\s+(.*\S)\s*$")
RE_NUM = re.compile(r"^\s*(\d+)[.)]\s+(.*\S)\s*$")
RE_BULLET = re.compile(r"^\s*[—\-\*]\s+(.*\S)\s*$")
RE_SUBHEAD = re.compile(r"^\s*([А-Д])\.\s+(.*\S)\s*$")     # подраздел довода А/Б/В/Г/Д
RE_QUOTE = re.compile(r"^\s*>\s?(.*)$")
RE_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
RE_TABLE_SEP = re.compile(r"^\s*\|?[\s:\-|]+\|?\s*$")

MINIHEADS = ("Обжалуемые судебные акты", "Суть спора", "Участники спора")
PROSBA_MARKERS = ("ПРОШУ", "Требую", "Требуем", "Ходатайства", "Ходатайство")
APPENDIX_MARKER = "Приложение"
SIGN_MARKER = "Представитель"
# «хвост» просительной, вводящий второй список (расширенный вариант жалобы)
RE_PROSBA_TAIL = re.compile(r"^\s*(При новом рассмотрении.*|.*просит суд:)\s*$")

TYPE_MAP = {
    "appeal": "жалоба", "cassation": "жалоба", "nadzor": "жалоба",
    "жалоба": "жалоба", "апелляционная": "жалоба", "кассационная": "жалоба",
    "otzyv": "иск-документ", "отзыв": "иск-документ",
    "hodataystvo": "иск-документ", "ходатайство": "иск-документ",
    "zayavlenie": "иск-документ", "заявление": "иск-документ",
    "vozrazhenie": "иск-документ", "возражение": "иск-документ",
    "poyasneniya": "иск-документ", "пояснения": "иск-документ",
    "pretenziya": "претензия", "претензия": "претензия",
}


# --------------------------------------------------------------------------- #
# OOXML-хелперы
# --------------------------------------------------------------------------- #
def _set_font(run, size, bold=False, name=FONT, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _no_char_spacing(run):
    """Сброс разрядки (style-spec §0 п.2): character spacing = 0."""
    rpr = run._element.get_or_add_rPr()
    sp = rpr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr.append(sp)
    sp.set(qn("w:val"), "0")


def _run_shade(run, fill):
    """Заливка за текстом (highlight #D9D9D9 через w:shd в rPr)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def _para_border(p, sides):
    """sides: dict, напр. {'bottom': True} (Heading2) или {'left': True} (цитата)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    for side in sides:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")       # 8 восьмых пункта = 1pt
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), "000000")
        pbdr.append(el)


def _cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _para(p, before=0, after=8, align=WD_ALIGN_PARAGRAPH.LEFT,
          left_indent=None, hanging=None, keep_with_next=False, line=True):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.alignment = align
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = LINE_MULT
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.first_line_indent = Pt(-hanging.pt if hasattr(hanging, "pt") else -hanging)
    if keep_with_next:
        pf.keep_with_next = True
    return p


# --- нумерация (style-spec §8 главный риск-узел): рестарт с 1 на каждый блок --- #
class Numbering:
    """Каждый вызов new_list() создаёт свой abstractNum+num (start=1) -> рестарт."""
    def __init__(self, doc):
        self._ok = True
        try:
            self._numbering = doc.part.numbering_part.element
        except Exception:  # numbering-часть отсутствует -> деградация в текстовые маркеры
            self._ok = False
            self._numbering = None
        self._abs = 900
        self._num = 900

    def available(self):
        return self._ok

    def new_list(self, is_bullet=False, ind_left=720, ind_hanging=360,
                 num_bold=False, num_size=None):
        self._abs += 1
        self._num += 1
        an = OxmlElement("w:abstractNum")
        an.set(qn("w:abstractNumId"), str(self._abs))
        mlt = OxmlElement("w:multiLevelType")
        mlt.set(qn("w:val"), "singleLevel")
        an.append(mlt)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "bullet" if is_bullet else "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "—" if is_bullet else "%1.")  # em-dash / «1.»
        lvl.append(lvlText)
        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        lvl.append(lvlJc)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(ind_left))
        ind.set(qn("w:hanging"), str(ind_hanging))
        ppr.append(ind)
        lvl.append(ppr)
        # rPr номера/маркера: Garamond всегда; bold и размер — по запросу (правка 3)
        rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), FONT)
        rf.set(qn("w:hAnsi"), FONT)
        rf.set(qn("w:cs"), FONT)
        rpr.append(rf)
        if num_bold:
            rpr.append(OxmlElement("w:b"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int((num_size or SZ_BODY) * 2)))
        rpr.append(sz)
        lvl.append(rpr)
        an.append(lvl)
        # все w:abstractNum должны идти ДО всех w:num
        first_num = self._numbering.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(an)
        else:
            self._numbering.append(an)
        n = OxmlElement("w:num")
        n.set(qn("w:numId"), str(self._num))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(self._abs))
        n.append(aid)
        self._numbering.append(n)
        return self._num

    @staticmethod
    def apply(p, num_id, ilvl=0):
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        il = OxmlElement("w:ilvl")
        il.set(qn("w:val"), str(ilvl))
        numPr.append(il)
        ni = OxmlElement("w:numId")
        ni.set(qn("w:val"), str(num_id))
        numPr.append(ni)
        pPr.append(numPr)


# --------------------------------------------------------------------------- #
# Парсер .md по якорному контракту -> модель блоков
# --------------------------------------------------------------------------- #
def parse_md(text):
    """Возвращает (header_lines, blocks). blocks — список dict со ключом 'type'."""
    raw = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    # 1) шапка = всё до первого H1
    h1_idx = next((i for i, l in enumerate(raw) if RE_H1.match(l)), None)
    header_lines, body_start = [], 0
    if h1_idx is not None:
        header_lines = [l for l in raw[:h1_idx] if l.strip()]
        body_start = h1_idx
    lines = raw[body_start:]

    blocks = []
    i, n = 0, len(lines)
    pending_num, pending_bullet = [], []

    def flush_lists():
        nonlocal pending_num, pending_bullet
        if pending_num:
            blocks.append({"type": "numlist", "items": pending_num})
            pending_num = []
        if pending_bullet:
            blocks.append({"type": "bullet", "items": pending_bullet})
            pending_bullet = []

    while i < n:
        line = lines[i]
        s = line.strip()
        if not s:
            flush_lists()
            i += 1
            continue

        m = RE_H1.match(line)
        if m:
            flush_lists()
            blocks.append({"type": "title", "text": m.group(1)})
            i += 1
            continue
        m = RE_DATE.match(line)
        if m and blocks and blocks[-1]["type"] == "title":
            blocks.append({"type": "date", "text": m.group(1)})
            i += 1
            continue
        m = RE_H2.match(line)
        if m:
            flush_lists()
            blocks.append({"type": "h2", "text": m.group(1)})
            i += 1
            continue
        # таблица
        if RE_TABLE_ROW.match(line):
            flush_lists()
            rows, j = [], i
            while j < n and RE_TABLE_ROW.match(lines[j]):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                if not RE_TABLE_SEP.match(lines[j]):
                    rows.append(cells)
                j += 1
            blocks.append({"type": "table", "rows": rows})
            i = j
            continue
        # цитата
        m = RE_QUOTE.match(line)
        if m:
            flush_lists()
            qlines, j = [], i
            while j < n and RE_QUOTE.match(lines[j]):
                qlines.append(RE_QUOTE.match(lines[j]).group(1).strip())
                j += 1
            blocks.append({"type": "quote", "text": " ".join(x for x in qlines if x)})
            i = j
            continue
        # маркеры-заголовки просительной / приложений / подписи
        if s.startswith(APPENDIX_MARKER):
            flush_lists()
            blocks.append({"type": "appendix_head", "text": s})
            i += 1
            continue
        if s.startswith(SIGN_MARKER):
            flush_lists()
            role, name = s, ""
            if "\t" in s:
                role, name = s.split("\t", 1)
            else:
                parts = s.rsplit("  ", 1)
                if len(parts) == 2:
                    role, name = parts[0].strip(), parts[1].strip()
            blocks.append({"type": "signature", "role": role.strip(), "name": name.strip()})
            i += 1
            continue
        if any(s.rstrip(":").startswith(pm) for pm in PROSBA_MARKERS):
            flush_lists()
            blocks.append({"type": "prosba_head", "text": s})
            i += 1
            continue
        if RE_PROSBA_TAIL.match(line):
            flush_lists()
            blocks.append({"type": "prosba_tail", "text": s})
            i += 1
            continue
        # мини-разделы жалобы
        if any(s.rstrip(":").startswith(mh) for mh in MINIHEADS):
            flush_lists()
            blocks.append({"type": "minihead", "text": s})
            i += 1
            continue
        # подраздел довода А/Б/В
        m = RE_SUBHEAD.match(line)
        if m and len(s) <= 90:
            flush_lists()
            blocks.append({"type": "subhead", "text": s})
            i += 1
            continue
        # нумерованный / маркированный список
        m = RE_NUM.match(line)
        if m:
            pending_num.append(m.group(2))
            i += 1
            continue
        m = RE_BULLET.match(line)
        if m:
            pending_bullet.append(m.group(1))
            i += 1
            continue
        # обычный абзац (возможная ссылка-источник после цитаты)
        flush_lists()
        if blocks and blocks[-1]["type"] == "quote" and _looks_like_source(s):
            blocks.append({"type": "source", "text": s})
        else:
            blocks.append({"type": "para", "text": s})
        i += 1

    flush_lists()
    return header_lines, blocks


def _looks_like_source(s):
    return bool(re.match(
        r"^(Постановление|Определение|Решение|Обзор|Пленум|п\.|ст\.|абз\.)", s
    )) and len(s) <= 160


# --------------------------------------------------------------------------- #
# Рендер модели -> .docx
# --------------------------------------------------------------------------- #
def _add_header_para(doc, lines, space_after=0):
    """Один абзац шапки из списка строк (внутри — soft return); leftIndent 210."""
    p = doc.add_paragraph()
    _para(p, before=0, after=space_after, left_indent=HEADER_LEFT_INDENT)
    for k, part in enumerate(lines):
        if k:
            p.add_run().add_break(WD_BREAK.LINE)
        run = p.add_run(part.strip())
        _set_font(run, SZ_BODY)
        _no_char_spacing(run)
    return p


def _group_header(header_lines):
    """Сгруппировать строки шапки в логические абзацы (пустая строка = разделитель).
    В плоском списке header_lines пустые уже отфильтрованы, поэтому группируем
    по маркерам «В »/«Дело »/«от »/«Кому:»/«От:»/«Исх.»."""
    groups, cur = [], []
    starters = ("В ", "Дело", "от ", "От:", "Кому:", "Исх.")
    for l in header_lines:
        st = l.strip()
        if cur and any(st.startswith(x) for x in starters):
            groups.append(cur)
            cur = [st]
        else:
            cur.append(st)
    if cur:
        groups.append(cur)
    return groups


def render(header_lines, blocks, out_path, case=None):
    doc = Document()
    # базовый стиль Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BODY)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = LINE_MULT

    # страница A4 + поля (рос. практика; левое шире под подшивку) + автоперенос (правка 1)
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Mm(297), Mm(210)
    sec.left_margin, sec.right_margin = Cm(3), Cm(1.5)
    sec.top_margin, sec.bottom_margin = Cm(2), Cm(2)
    _enable_hyphenation(doc)

    numbering = Numbering(doc)
    _H2_STATE["num_id"] = None   # сброс сквозной нумерации доводов на каждый рендер

    # --- шапка (правка 2: интервал после «Дело №…» перед «от …») ---
    for group in _group_header(header_lines):
        sa = 8 if group and group[0].strip().startswith("Дело") else 0
        _add_header_para(doc, group, space_after=sa)

    for idx, b in enumerate(blocks):
        t = b["type"]

        if t == "title":
            p = doc.add_paragraph()
            _para(p, before=24, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_H1, bold=True)

        elif t == "date":
            p = doc.add_paragraph()
            _para(p, before=0, after=8)
            r = p.add_run(b["text"])
            _set_font(r, SZ_BODY)

        elif t == "minihead":
            p = doc.add_paragraph()
            _para(p, before=8, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_MINIHEAD, bold=True)

        elif t == "h2":
            p = doc.add_paragraph()
            _para(p, before=16, after=8, keep_with_next=True)
            _para_border(p, ["bottom"])
            r = p.add_run(b["text"])   # номер проставит автонумерация (см. ниже)
            _set_font(r, SZ_H2, bold=True)
            # привязка Heading2-довода к сквозному списку доводов (автонумерация 1, 2, 3…)
            _apply_h2_number(doc, p, numbering)

        elif t == "subhead":
            p = doc.add_paragraph()
            _para(p, before=8, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_SUBHEAD, bold=True)

        elif t == "prosba_head":
            p = doc.add_paragraph()
            _para(p, before=16, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_PROSBA, bold=True)

        elif t == "prosba_tail":
            p = doc.add_paragraph()
            _para(p, before=8, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_BODY)

        elif t == "appendix_head":
            p = doc.add_paragraph()
            _para(p, before=8, after=0)
            r = p.add_run(b["text"])
            _set_font(r, SZ_BODY, bold=True)

        elif t == "para":
            p = doc.add_paragraph()
            _para(p, before=0, after=8)
            _emit_text_with_highlight(p, b["text"], in_argument=_in_argument(blocks, idx))

        elif t == "quote":
            p = doc.add_paragraph()
            _para(p, before=0, after=0, left_indent=QUOTE_INDENT)  # правка 7: блок без интервала снизу
            _para_border(p, ["left"])
            r = p.add_run(b["text"])
            _set_font(r, SZ_BODY)

        elif t == "source":
            p = doc.add_paragraph()
            _para(p, before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT,
                  left_indent=QUOTE_INDENT)  # правка 7: интервал даст следующий текст (spaceBefore)
            _para_border(p, ["left"])
            r = p.add_run(b["text"])
            _set_font(r, SZ_BODY)

        elif t == "numlist":
            _emit_list(doc, numbering, b["items"], is_bullet=False)

        elif t == "bullet":
            _emit_list(doc, numbering, b["items"], is_bullet=True)

        elif t == "table":
            _emit_table(doc, b["rows"])

        elif t == "signature":
            p = doc.add_paragraph()
            _para(p, before=24, after=0)   # больше воздуха сверху под подпись
            # правка 8: ФИО к правому краю, широкий зазор между ролью и ФИО под подпись
            content_w = sec.page_width - sec.left_margin - sec.right_margin
            p.paragraph_format.tab_stops.add_tab_stop(content_w, WD_TAB_ALIGNMENT.RIGHT)
            r = p.add_run(b["role"] + ("\t" + b["name"] if b["name"] else ""))
            _set_font(r, SZ_BODY, bold=True)

    _fix_after_table_spacing(doc)
    _fix_after_citation_spacing(doc)
    _set_deterministic_metadata(doc)
    doc.save(out_path)


# доводы: единый сквозной список (1, 2, 3…) на все Heading2
_H2_STATE = {"num_id": None}


def _apply_h2_number(doc, p, numbering):
    if not numbering.available():
        return
    if _H2_STATE["num_id"] is None:
        # правки 3, 4: номер довода жирный 14pt и у левого поля (не индентирован)
        _H2_STATE["num_id"] = numbering.new_list(
            is_bullet=False, ind_left=397, ind_hanging=397,
            num_bold=True, num_size=SZ_H2)
    Numbering.apply(p, _H2_STATE["num_id"])


def _emit_list(doc, numbering, items, is_bullet):
    if numbering.available():
        num_id = numbering.new_list(is_bullet=is_bullet)  # рестарт с 1 на каждый блок
        for k, it in enumerate(items):
            p = doc.add_paragraph()
            last = (k == len(items) - 1)
            _para(p, before=0, after=(12 if last else 0),
                  left_indent=LIST_LEFT, hanging=LIST_HANG)
            Numbering.apply(p, num_id)
            r = p.add_run(it)
            _set_font(r, SZ_BODY)
    else:
        # деградация: текстовые маркеры (нумерация недоступна)
        for k, it in enumerate(items):
            p = doc.add_paragraph()
            last = (k == len(items) - 1)
            _para(p, before=0, after=(12 if last else 0), left_indent=LIST_LEFT)
            prefix = "— " if is_bullet else "%d. " % (k + 1)
            r = p.add_run(prefix + it)
            _set_font(r, SZ_BODY)


def _emit_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.LEFT   # правка 5: на одном уровне с текстом
    table.autofit = True
    _table_autofit_layout(table)                # правка 6: автоподбор ширины по содержимому
    for ri, row in enumerate(rows):
        is_header = (ri == 0)
        is_total = _is_total_row(row)
        for ci in range(ncol):
            cell = table.rows[ri].cells[ci]
            txt = row[ci] if ci < len(row) else ""
            para = cell.paragraphs[0]      # свежая ячейка: пустой абзац без runs
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.0   # ячейки таблиц — single (style-spec §1)
            # выравнивание: числовые -> Right
            if _is_numeric(txt) and not is_header:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = para.add_run(txt)
            _set_font(r, SZ_TABLE, bold=(is_header or is_total), name=TABLE_FONT)
            if is_header:
                _cell_shade(cell, HL_FILL)
            elif is_total:
                _cell_shade(cell, TOTAL_FILL)


def _emit_text_with_highlight(p, text, in_argument):
    """Обычный абзац; внутри абзаца довода выделяем одну ключевую фразу серым."""
    r_all = None
    if in_argument:
        phrase = _key_phrase(text)
        if phrase:
            before, _, after = text.partition(phrase)
            if before:
                _set_font(p.add_run(before), SZ_BODY)
            r = p.add_run(phrase)
            _set_font(r, SZ_BODY)
            _run_shade(r, HL_FILL)
            if after:
                _set_font(p.add_run(after), SZ_BODY)
            return
    r_all = p.add_run(text)
    _set_font(r_all, SZ_BODY)


# --------------------------------------------------------------------------- #
# Вспомогательное
# --------------------------------------------------------------------------- #
KEY_MARKERS = (
    "подлежит отмене", "не соответствует материалам дела", "не соответствует",
    "противоречит материалам дела", "не носит абстрактного характера",
    "неверно применил", "ошибочно", "нарушил", "подлежит удовлетворению",
    "не подлежит",
)


def _key_phrase(text):
    """Найти ключевую фразу для highlight (модальные конструкции). Возвращает
    подстроку-предложение с маркером или None (тогда highlight не ставится)."""
    for m in KEY_MARKERS:
        idx = text.find(m)
        if idx != -1:
            start = text.rfind(".", 0, idx)
            start = start + 1 if start != -1 else 0
            end = text.find(".", idx)
            end = end + 1 if end != -1 else len(text)
            return text[start:end].strip()
    return None


def _in_argument(blocks, idx):
    """Абзац находится внутри раздела-довода (после ближайшего h2 до следующего)?"""
    for j in range(idx, -1, -1):
        t = blocks[j]["type"]
        if t == "h2":
            return True
        if t in ("prosba_head", "title", "minihead", "appendix_head"):
            return False
    return False


def _is_numeric(s):
    s2 = re.sub(r"[   ]", "", s).replace(",", "").replace("%", "")
    s2 = s2.replace("—", "").replace("-", "").replace(".", "")
    return s2.isdigit() and any(ch.isdigit() for ch in s)


def _is_total_row(row):
    return any(c.strip().rstrip(":").lower() in ("итого", "всего") for c in row)


def _fix_after_table_spacing(doc):
    """style-spec §0 п.14 / §3: абзац сразу после таблицы -> spaceBefore 16pt."""
    body = doc.element.body
    children = list(body)
    for k, el in enumerate(children):
        if el.tag == qn("w:tbl") and k + 1 < len(children):
            nxt = children[k + 1]
            if nxt.tag == qn("w:p"):
                pPr = nxt.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    nxt.insert(0, pPr)
                spacing = pPr.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    pPr.append(spacing)
                spacing.set(qn("w:before"), "320")   # 16pt = 320 twips
                spacing.set(qn("w:line"), "276")      # 1.15*240
                spacing.set(qn("w:lineRule"), "auto")


def _fix_after_citation_spacing(doc):
    """Правка 7: после цитатного блока (абзацы с левой границей) — интервал сверху
    у следующего обычного текста; сам блок цитаты — без интервала снизу."""
    def has_left_border(p):
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            return False
        pbdr = pPr.find(qn("w:pBdr"))
        return pbdr is not None and pbdr.find(qn("w:left")) is not None
    paras = doc.paragraphs
    for i in range(len(paras) - 1):
        if has_left_border(paras[i]) and not has_left_border(paras[i + 1]):
            pf = paras[i + 1].paragraph_format
            sb = pf.space_before
            if sb is None or sb < Pt(16):
                pf.space_before = Pt(16)


def _enable_hyphenation(doc):
    """Правка 1: автоматический перенос слов на уровне документа."""
    settings = doc.settings.element
    if settings.find(qn("w:autoHyphenation")) is None:
        el = OxmlElement("w:autoHyphenation")
        el.set(qn("w:val"), "true")
        settings.insert(0, el)


# порядок дочерних элементов w:tblPr по ECMA-376 (иначе Word игнорирует свойства)
_TBLPR_ORDER = [
    "w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
    "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
    "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd", "w:tblLayout",
    "w:tblCellMar", "w:tblLook", "w:tblCaption", "w:tblDescription",
]


def _tblpr_set(tblPr, tag, attrs):
    """Найти/создать дочерний элемент tblPr и вставить в правильную по схеме позицию."""
    el = tblPr.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        idx = _TBLPR_ORDER.index(tag)
        after = None
        for ex in tblPr:
            extag = "w:" + ex.tag.split("}")[-1]
            if extag in _TBLPR_ORDER and _TBLPR_ORDER.index(extag) > idx:
                after = ex
                break
        if after is not None:
            after.addprevious(el)
        else:
            tblPr.append(el)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el


def _table_autofit_layout(table):
    """Левая граница таблицы на уровне текста (tblInd=0 в правильном порядке tblPr) +
    автоподбор ширины по содержимому (AutoFit to Contents: tblLayout=autofit, tblW=auto,
    сняты фиксированные ширины колонок сетки и ячеек)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Word позиционирует таблицу по КРАЮ ТЕКСТА первой ячейки, а не по границе:
    # при tblInd=0 граница уходит влево от текста на левое поле ячейки. Ставим
    # tblInd = левому полю ячейки → граница садится ровно на поле полосы (уровень текста).
    cell_left = 108   # twips ≈ 0.19 см (стандартное левое поле ячейки Word)
    _tblpr_set(tblPr, "w:tblW", {"w:w": "0", "w:type": "auto"})
    _tblpr_set(tblPr, "w:tblInd", {"w:w": str(cell_left), "w:type": "dxa"})
    _tblpr_set(tblPr, "w:tblLayout", {"w:type": "autofit"})
    # зафиксировать поля ячеек, чтобы компенсация tblInd была точной
    cm = _tblpr_set(tblPr, "w:tblCellMar", {})
    for ch in list(cm):
        cm.remove(ch)
    for side, w in (("top", "0"), ("left", str(cell_left)),
                    ("bottom", "0"), ("right", str(cell_left))):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), w)
        m.set(qn("w:type"), "dxa")
        cm.append(m)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in grid.findall(qn("w:gridCol")):
            if gc.get(qn("w:w")) is not None:
                del gc.attrib[qn("w:w")]
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcW.set(qn("w:type"), "auto")
                tcW.set(qn("w:w"), "0")


def _set_deterministic_metadata(doc):
    cp = doc.core_properties
    cp.author = "vassal-litigator"
    cp.last_modified_by = "vassal-litigator"
    fixed = datetime(2000, 1, 1, 0, 0, 0)
    cp.created = fixed
    cp.modified = fixed
    cp.revision = 1


# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="Оформление процессуального документа .md -> .docx")
    ap.add_argument("in_md")
    ap.add_argument("out_docx")
    ap.add_argument("--type", default="", help="тип документа (см. докстринг)")
    ap.add_argument("--case", default=None, help="путь к case.yaml (для сборки шапки, опц.)")
    args = ap.parse_args()

    doc_type = TYPE_MAP.get(args.type.strip().lower(), args.type.strip() or "иск-документ")

    with open(args.in_md, encoding="utf-8") as f:
        text = f.read()

    header_lines, blocks = parse_md(text)

    case = None
    if args.case:
        try:
            import yaml
            with open(args.case, encoding="utf-8") as f:
                case = yaml.safe_load(f)
        except Exception as e:
            sys.stderr.write("format_doc: case.yaml не прочитан (%s) — шапка из .md.\n" % e)

    # если в .md шапки нет, а case есть — минимальная сборка (best-effort)
    if not header_lines and case:
        header_lines = _header_from_case(case, doc_type)

    render(header_lines, blocks, args.out_docx, case=case)
    sys.stderr.write("format_doc: готово -> %s (тип: %s, блоков: %d)\n"
                     % (args.out_docx, doc_type, len(blocks)))


def _header_from_case(case, doc_type):
    c = case.get("case", case) if isinstance(case, dict) else {}
    court = c.get("court", "")
    number = c.get("number", "")
    lines = []
    if court:
        lines.append("В %s" % court)
    if number:
        lines.append("Дело №%s" % number)
    return [l for l in lines if l]


if __name__ == "__main__":
    main()
